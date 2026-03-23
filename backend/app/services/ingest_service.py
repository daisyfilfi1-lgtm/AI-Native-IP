"""
素材录入流水线：根据任务拉取内容、分块、写入 ip_assets，并更新 ingest_tasks。
支持：text/document 的 URL 或本地上传（UTF-8 的 txt/md、docx 正文、PDF 可选字层文本）；video/audio 的 URL 或本地上传（Whisper 转写）。
已配置 LLM 时，会自动调用打标（asset_meta.auto_labels）。

防OOM优化：
- 任务超时控制（软超时，定期检查）
- 内存使用限制
- 更频繁的数据库提交
"""
import gc
import concurrent.futures
import io
import logging
import os
import tempfile
import time
import uuid
import zipfile
from pathlib import Path
from typing import List, Optional

import requests
from sqlalchemy.orm import Session

from app.db.models import FileObject, IPAsset, IngestTask, TagConfig
from app.services.ai_client import (
    embed_texts_batched,
    suggest_tags_for_content,
    transcribe,
    transcribe_from_url,
)
from app.db.session import SessionLocal
from app.services.storage_service import download_bytes
from app.services.vector_service import upsert_asset_vector

# 分块默认长度（字符），可按需调整
CHUNK_SIZE = 2000
CHUNK_OVERLAP = 100

# 大素材保护：避免海量分块 ×（每块 LLM + Embedding）拖垮内存与 CPU（OOM / Killed）
def _ingest_max_text_chars() -> int:
    raw = os.environ.get("INGEST_MAX_TEXT_CHARS", "").strip()
    if not raw:
        return 500_000
    try:
        return max(10_000, int(raw))
    except ValueError:
        return 500_000


def _ingest_max_chunks() -> int:
    raw = os.environ.get("INGEST_MAX_CHUNKS", "").strip()
    if not raw:
        return 120
    try:
        return max(1, int(raw))
    except ValueError:
        return 120


def _ingest_embed_batch_size() -> int:
    raw = os.environ.get("INGEST_EMBED_BATCH_SIZE", "").strip()
    if not raw:
        return 16
    try:
        return max(1, min(64, int(raw)))
    except ValueError:
        return 16


def _ingest_commit_every() -> int:
    raw = os.environ.get("INGEST_COMMIT_EVERY", "").strip()
    if not raw:
        return 25
    try:
        return max(1, int(raw))
    except ValueError:
        return 25


def _ingest_max_url_bytes() -> int:
    raw = os.environ.get("INGEST_MAX_URL_BYTES", "").strip()
    if not raw:
        return 20 * 1024 * 1024
    try:
        return max(64 * 1024, int(raw))
    except ValueError:
        return 20 * 1024 * 1024


def _ingest_skip_embedding() -> bool:
    """整段录入不写向量（仅文本入库），用于极低内存或排查 OOM。"""
    return os.environ.get("INGEST_SKIP_EMBEDDING", "").lower() in ("1", "true", "yes")


def _ingest_task_timeout() -> int:
    """任务执行超时时间（秒），防止Railway Kill进程。"""
    raw = os.environ.get("INGEST_TASK_TIMEOUT", "").strip()
    if not raw:
        return 25  # 默认25秒安全值
    try:
        return max(10, min(60, int(raw)))  # 限制10-60秒
    except ValueError:
        return 25


class TimeoutException(Exception):
    """任务超时异常"""
    pass


class TimeLimitChecker:
    """软超时检查器（线程安全，适用于 Railway 后台任务）"""
    def __init__(self, seconds: int):
        self.seconds = seconds
        self.start_time = time.time()
        self.timed_out = False
    
    def check(self):
        """检查是否已超时，如果超时则抛出异常"""
        if time.time() - self.start_time > self.seconds:
            self.timed_out = True
            raise TimeoutException(f"Task exceeded {self.seconds} seconds")
    
    def remaining(self) -> float:
        """返回剩余时间"""
        return max(0, self.seconds - (time.time() - self.start_time))


def _call_with_timeout(
    fn,
    *args,
    timeout_seconds: float,
    **kwargs,
):
    """
    在独立线程执行潜在阻塞调用（LLM/Embedding），避免长时间卡死任务状态。
    """
    if timeout_seconds <= 0:
        raise TimeoutException("Task exceeded timeout while waiting external service")
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(fn, *args, **kwargs)
        try:
            return future.result(timeout=timeout_seconds)
        except concurrent.futures.TimeoutError as e:
            raise TimeoutException(
                f"External service call timed out after {timeout_seconds:.1f}s"
            ) from e


def get_ingest_task(db: Session, task_id: str) -> IngestTask | None:
    return db.query(IngestTask).filter(IngestTask.task_id == task_id).first()


def _normalize_http_url(url: str | None) -> str:
    """
    补全 http(s) 协议，避免 requests / urllib 抛出 MissingSchema。
    用户常漏写协议（如只填域名或误填「ddd」）。
    """
    u = (url or "").strip()
    if not u:
        return ""
    if u.startswith("//"):
        return "https:" + u
    low = u.lower()
    if low.startswith("http://") or low.startswith("https://"):
        return u
    return "https://" + u


def _fetch_text_from_url(url: str) -> str:
    """流式拉取 URL，避免超大响应一次性占满内存。"""
    resolved = _normalize_http_url(url)
    if not resolved:
        raise ValueError("source_url 为空或仅空白，请填写以 http:// 或 https:// 开头的完整链接。")
    max_bytes = _ingest_max_url_bytes()
    try:
        resp = requests.get(resolved, timeout=60, stream=True)
    except requests.RequestException as e:
        raise ValueError(
            f"无法拉取 URL（{resolved}）：{e}"
        ) from e
    try:
        resp.raise_for_status()
        buf = bytearray()
        for chunk in resp.iter_content(chunk_size=64 * 1024):
            if not chunk:
                continue
            buf.extend(chunk)
            if len(buf) > max_bytes:
                raise ValueError(
                    f"URL 正文超过 INGEST_MAX_URL_BYTES 上限（{max_bytes} bytes）"
                )
        enc = (resp.encoding or "utf-8").strip() or "utf-8"
        return bytes(buf).decode(enc, errors="replace")
    finally:
        resp.close()


def _transcribe_from_file_object(db: Session, task: IngestTask) -> str:
    """从对象存储已上传文件转写音视频（Whisper）。"""
    if not task.local_file_id:
        return ""
    row = (
        db.query(FileObject)
        .filter(FileObject.file_id == task.local_file_id, FileObject.ip_id == task.ip_id)
        .first()
    )
    if not row:
        return ""
    data = download_bytes(row.bucket, row.object_key)
    if not data:
        return ""
    suffix = Path(row.file_name or row.object_key or "").suffix or ".bin"
    tmp_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        return (transcribe(tmp_path) or "").strip()
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _text_from_docx(data: bytes) -> str:
    """从 .docx（OOXML）提取纯文本：段落 + 表格单元格。"""
    try:
        from docx import Document
    except ImportError:
        return ""
    try:
        doc = Document(io.BytesIO(data))
        parts: List[str] = []
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        parts.append(t)
        return "\n".join(parts)
    except Exception:
        return ""


def _looks_like_pdf(data: bytes) -> bool:
    return len(data) >= 5 and data[:5] == b"%PDF-"


def _looks_like_docx_ooxml(data: bytes) -> bool:
    """docx 为 ZIP，内含 word/document.xml（与 xlsx/pptx 区分）。"""
    if len(data) < 4 or data[:2] != b"PK":
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            names = zf.namelist()
            return any(n == "word/document.xml" or n.endswith("/word/document.xml") for n in names)
    except Exception:
        return False


def _text_from_pdf(data: bytes) -> str:
    """
    从 PDF 提取可选中文字层（非扫描件）。扫描版 PDF 无文本层时结果为空，需 OCR 另处理。
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(io.BytesIO(data))
        parts: List[str] = []
        for page in reader.pages:
            t = (page.extract_text() or "").strip()
            if t:
                parts.append(t)
        return "\n".join(parts)
    except Exception:
        return ""


def _load_text_from_file_object(db: Session, task: IngestTask) -> str:
    if not task.local_file_id:
        return ""
    row = (
        db.query(FileObject)
        .filter(FileObject.file_id == task.local_file_id, FileObject.ip_id == task.ip_id)
        .first()
    )
    if not row:
        return ""
    data = download_bytes(row.bucket, row.object_key)
    if not data:
        return ""
    try:
        name = (row.file_name or row.object_key or "").lower()
        if name.endswith(".docx"):
            text = _text_from_docx(data)
            if text.strip():
                return text
            return ""
        if name.endswith(".doc"):
            # 旧版 Word 二进制 .doc，需 LibreOffice/专用解析；此处不冒充文本
            return ""
        if name.endswith(".pdf"):
            text = _text_from_pdf(data)
            if text.strip():
                return text
            return ""
        # 扩展名缺失或误命名时，用魔数补救（如 object_key 为 upload.bin）
        if _looks_like_pdf(data):
            text = _text_from_pdf(data)
            if text.strip():
                return text
        if _looks_like_docx_ooxml(data):
            text = _text_from_docx(data)
            if text.strip():
                return text
        # .txt / .md / .csv 等：按 UTF-8 读取
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return data.decode("utf-8", errors="ignore")
            except Exception:
                return ""
    finally:
        del data


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """按长度分块，带重叠，避免截断在句中对齐到换行。"""
    text = (text or "").strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            # 尽量在换行处截断
            last_nl = text.rfind("\n", start, end + 1)
            if last_nl > start:
                end = last_nl + 1
        chunks.append(text[start:end].strip())
        if not chunks[-1]:
            start = end
            continue
        start = end - overlap
        if start >= len(text):
            break
    return [c for c in chunks if c]


def _run_ingest_pipeline(
    db: Session,
    task: IngestTask,
    *,
    checker: Optional[TimeLimitChecker] = None,
) -> None:
    """执行录入：拉取内容 → 分块 → 写入 ip_assets → 更新 task。可选 checker 用于后台线程中的软超时（不可用 signal）。"""
    import logging

    logger = logging.getLogger(__name__)
    if checker:
        logger.info(
            "Starting pipeline for task: %s (soft timeout %ss)",
            task.task_id,
            checker.seconds,
        )
    else:
        logger.info("Starting pipeline for task: %s", task.task_id)

    task.status = "PROCESSING"
    db.commit()

    try:
        if checker:
            checker.check()
        st = task.source_type or "text"
        if st in ("video", "audio") and task.local_file_id:
            raw_text = _transcribe_from_file_object(db, task)
            if not (raw_text or "").strip():
                raw_text = (
                    f"[{st} 转写失败或未配置 Whisper] {task.title or task.local_file_id}\n\n"
                    "请检查 OPENAI_TRANSCRIPTION_API_KEY（或主 OPENAI_API_KEY）及文件格式。"
                )
        elif st in ("video", "audio") and task.source_url:
            src = _normalize_http_url(task.source_url)
            if not src:
                raw_text = (
                    f"[{st} 转写失败] source_url 无效或为空。\n\n"
                    "请填写以 http:// 或 https:// 开头的完整音视频链接。"
                )
            else:
                raw_text = transcribe_from_url(src)
            if not (raw_text or "").strip():
                raw_text = (
                    f"[{st} 转写失败或未配置 Whisper] {task.title or task.source_url or '未命名'}\n\n"
                    "请检查转写 API 及网络。"
                )
        elif st in ("text", "document", "file") and task.local_file_id:
            raw_text = _load_text_from_file_object(db, task)
            if not raw_text.strip():
                raw_text = (
                    f"[文件解析失败] {task.title or task.local_file_id}\n\n"
                    "已支持：UTF-8 文本（txt/md）、docx、含文字层的 PDF（扫描件无文字层则为空）。"
                    "若仍为失败，请确认文件未加密；旧版 .doc 尚未接入。"
                )
        elif st in ("text", "document", "file") and task.source_url:
            raw_text = _fetch_text_from_url(task.source_url)
        elif task.local_file_id:
            raw_text = _load_text_from_file_object(db, task)
            if not (raw_text or "").strip():
                raw_text = _transcribe_from_file_object(db, task)
            if not (raw_text or "").strip():
                raw_text = (
                    f"[文件解析失败] {task.title or task.local_file_id}\n\n"
                    "请确认 source_type 为 text/document/file（文本）或 audio/video（音视频）；"
                    "文本类已支持 txt/md、docx、PDF（文字层）。"
                )
        elif st in ("video", "audio"):
            raw_text = (
                f"[{st} 待转写] {task.title or '未命名'}\n\n"
                "请填写 source_url，或先上传文件并传入 local_file_id。"
            )
        else:
            raw_text = f"[待处理] {task.title or '未命名'}\n\nsource_type={st}"

        if checker:
            checker.check()

        max_chars = _ingest_max_text_chars()
        max_chunks = _ingest_max_chunks()
        embed_bs = _ingest_embed_batch_size()
        commit_every = _ingest_commit_every()
        skip_emb = _ingest_skip_embedding()

        ingest_notes: List[str] = []
        rt = raw_text or ""
        if len(rt) > max_chars:
            ingest_notes.append(f"正文已截断至前 {max_chars} 字符，避免录入任务 OOM")
            rt = rt[:max_chars]

        if skip_emb:
            ingest_notes.append("INGEST_SKIP_EMBEDDING=1：已跳过向量写入，仅保存文本与素材记录")

        chunks = _chunk_text(rt)
        if not chunks:
            chunks = [rt or "(无内容)"]

        if len(chunks) > max_chunks:
            ingest_notes.append(f"分块仅保留前 {max_chunks} 段（共 {len(chunks)} 段），请拆分素材或调大 INGEST_MAX_CHUNKS")
            chunks = chunks[:max_chunks]

        if checker:
            checker.check()

        tag_cfg = db.query(TagConfig).filter(TagConfig.ip_id == task.ip_id).first()
        categories = (tag_cfg.tag_categories if tag_cfg else None) or None

        if checker:
            checker.check()

        # 每块各打一次 LLM 会极慢且易爆内存；整任务只打标一次（用开头片段）
        shared_tags = None
        try:
            tag_snippet = (rt or "")[:4000]
            if tag_snippet.strip():
                tag_timeout = min(15.0, checker.remaining()) if checker else 15.0
                shared_tags = _call_with_timeout(
                    suggest_tags_for_content,
                    tag_snippet,
                    categories,
                    timeout_seconds=tag_timeout,
                )
        except TimeoutException:
            raise
        except Exception:
            shared_tags = None

        # 按批向量化并落库：每批仅保留当前批的向量，避免「全部分块 embedding 列表」撑爆内存（OOM / Killed）
        created_ids: List[str] = []
        total = len(chunks)
        for batch_start in range(0, total, embed_bs):
            if checker:
                checker.check()
            batch_chunks = chunks[batch_start : batch_start + embed_bs]
            if skip_emb:
                batch_embeddings: List[list[float] | None] = [None] * len(batch_chunks)
            else:
                embed_timeout = min(20.0, checker.remaining()) if checker else 20.0
                batch_embeddings = _call_with_timeout(
                    embed_texts_batched,
                    batch_chunks,
                    batch_size=min(embed_bs, len(batch_chunks)),
                    timeout_seconds=embed_timeout,
                )
            for j, content in enumerate(batch_chunks):
                i = batch_start + j
                asset_id = f"asset_{task.ip_id}_{uuid.uuid4().hex[:12]}"
                title = (task.title or "未命名") + (f" (片段{i+1})" if total > 1 else "")
                meta: dict = {
                    "source_task_id": task.task_id,
                    "source_type": task.source_type,
                    "chunk_index": i,
                    "total_chunks": total,
                }
                if ingest_notes and i == 0:
                    meta["ingest_warnings"] = ingest_notes
                if shared_tags:
                    meta["auto_labels"] = shared_tags
                db.add(
                    IPAsset(
                        asset_id=asset_id,
                        ip_id=task.ip_id,
                        asset_type="story",
                        title=title,
                        content=content,
                        content_vector_ref=None,
                        asset_meta=meta,
                        relations=[],
                        status="active",
                    )
                )
                try:
                    emb = batch_embeddings[j] if j < len(batch_embeddings) else None
                    upsert_asset_vector(
                        db,
                        asset_id=asset_id,
                        ip_id=task.ip_id,
                        content=content,
                        precomputed_embedding=emb,
                    )
                except Exception:
                    pass
                created_ids.append(asset_id)

                if (i + 1) % commit_every == 0:
                    db.commit()
                    db.refresh(task)

            del batch_embeddings

        task.status = "COMPLETED"
        task.created_asset_ids = created_ids
        task.error_message = None
    except TimeoutException:
        raise
    except Exception as e:
        logger.exception("ingest pipeline failed: %s", task.task_id)
        task.status = "FAILED"
        task.error_message = str(e)
        task.created_asset_ids = []

    db.commit()


def process_ingest_task(task_id: str) -> None:
    """
    后台执行录入任务（在独立会话中）。
    使用软超时机制（定期检查），适用于 Railway 后台任务线程。
    """
    logger = logging.getLogger(__name__)
    start_time = time.time()
    timeout_seconds = _ingest_task_timeout()
    
    logger.info(f"[Ingest] Starting task {task_id}, timeout={timeout_seconds}s")
    
    db = SessionLocal()
    try:
        task = get_ingest_task(db, task_id)
        if not task:
            logger.warning(f"[Ingest] Task {task_id} not found")
            return
        if task.status != "QUEUED":
            logger.warning(f"[Ingest] Task {task_id} already processed: {task.status}")
            return

        # 执行流水线（软超时，适用于后台线程，禁止使用 signal.SIGALRM）
        try:
            _run_ingest_pipeline(
                db,
                task,
                checker=TimeLimitChecker(timeout_seconds),
            )
            
            duration = time.time() - start_time
            logger.info(f"[Ingest] Completed task {task_id} in {duration:.1f}s")
            
        except TimeoutException:
            duration = time.time() - start_time
            logger.error(f"[Ingest] Task {task_id} TIMEOUT after {duration:.1f}s")
            
            db.refresh(task)
            task.status = "FAILED"
            task.error_message = (
                f"处理超时({timeout_seconds}秒)。"
                f"建议：1) 压缩文件 2) 拆分成小文件(<5MB) 3) 稍后重试"
            )
            db.commit()
            
    except Exception as e:
        duration = time.time() - start_time
        logger.exception(f"[Ingest] Task {task_id} FAILED after {duration:.1f}s: {e}")
        
        try:
            db.refresh(task)
            task.status = "FAILED"
            task.error_message = str(e)[:500]
            db.commit()
        except Exception as inner_e:
            logger.error(f"[Ingest] Failed to update error status: {inner_e}")
    finally:
        db.close()
        gc.collect()
        logger.info(f"[Ingest] Task {task_id} cleanup complete")
